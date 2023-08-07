from flask import *
import os
from flask_mysqldb import MySQL
import fitz
from docx import Document
import re
import docx
from PyPDF2 import PdfReader

app = Flask(__name__)

app.config['MYSQL_HOST'] = 'localhost'
app.config['MYSQL_USER'] = 'root'
app.config['MYSQL_PASSWORD'] = 'Aravind@431'
app.config['MYSQL_DB'] = 'HR'
app.config['MYSQL_CURSORCLASS'] = 'DictCursor'

mysql = MySQL(app)

def extract_docx(file):
    document = docx.Document(file)
    return document.paragraphs[0].text.strip()

def extract_pdf(file):
    pdf_reader = PdfReader(file)
    page = pdf_reader.pages[0]
    page_text = page.extract_text()
    return page_text.split('\n')[0].strip()

def extract_docx_Details(file):
    document = docx.Document(file)
    email, mobile = '', ''
    for para in document.paragraphs:
        if 'email' in para.text.lower():
            email = para.text.split(':')[-1].strip()
        if 'mobile' in para.text.lower() or 'phone' in para.text.lower() or 'Mobile No' in para.text.lower():
            mobile = para.text.split(':')[-1].strip()
    return email, mobile

def extract_pdf_Details(file):
    pdf_reader = PdfReader(file)
    email, mobile = '', ''
    for page_num in range(len(pdf_reader.pages)):
        page = pdf_reader.pages[page_num]
        page_text = page.extract_text()
        if 'email' in page_text.lower():
            email = page_text.split('Email:')[-1].split('\n')[0].strip()
        if 'mobile' in page_text.lower() or 'phone' in page_text.lower():
            mobile_line = [line.strip() for line in page_text.split('\n') if 'mobile' in line.lower() or 'phone' in line.lower()]
            if mobile_line:
                mobile = mobile_line[0].split(':')[-1].strip()
            break  

    return email, mobile

def create_table():
    cur = mysql.connection.cursor()
    cur.execute('''CREATE TABLE IF NOT EXISTS skills_list(id INT AUTO_INCREMENT PRIMARY KEY,
                skills VARCHAR(255),
                name VARCHAR(250),
                FOREIGN KEY (name) REFERENCES candidates(name))''')
    mysql.connection.commit()
    cur.close()
    
def create_candidatesTable():
    cur = mysql.connection.cursor()
    cur.execute('''CREATE TABLE IF NOT EXISTS candidates 
                (id INT AUTO_INCREMENT PRIMARY KEY,
                name VARCHAR(255) UNIQUE,
                email VARCHAR(255),
                mobile_number varchar(255),
                resume LONGBLOB)''')
    mysql.connection.commit()
    cur.close()

@app.route('/')
def index():
    return render_template('login.html')

@app.route('/Home',methods=['POST'])
def home():
    un = request.form['uname']
    ps = request.form['passw']

    cur = mysql.connection.cursor()
    cur.execute("SELECT * FROM Userl_login WHERE username = %s AND password = %s", (un, ps))
    user = cur.fetchone()
    cur.close()

    if user:
        return render_template('index.html')
    else:
        error="Invalid credentials,Please enter valid credentials"
        return render_template('login.html',error1=error)

@app.route('/search', methods=['POST'])
def search_words():
    words_input = request.form.get('words')
    words = [word.strip().lower() for word in words_input.split(',')]
    file = request.files['file']
    filenm = file.filename

    if file and allowed_file(file.filename):
        if file.filename.lower().endswith('.pdf'):
            text = extract_text_from_pdf(file)
            name=extract_pdf(file)
            email,mobile=extract_pdf_Details(file)
        elif file.filename.lower().endswith('.docx'):
            text = extract_text_from_docx(file)
            name=extract_docx(file)
            email,mobile=extract_docx_Details(file)
        else:
            return jsonify(error='Unsupported file format. Only PDF and DOCX files are allowed.')

        #results = names + emails + mobile_numbers
        #print(results)
        matches = [word for word in words if word in text.lower()]
        cur = mysql.connection.cursor()
        create_candidatesTable()
        cur.execute('select * from candidates where email=%s',(email,))
        user=cur.fetchone()
        if user:
            error4="This profile is already exists"
            return render_template('index.html', msg=error4)
        else:
            res1 = len(matches)
            res2 = len(words)
            if res1==res2:
                sus="Profile shortlisted"
                store_candidateDetails(name,email,mobile)
                store_matches(matches,name)
                return render_template('index.html',sus=sus)
            elif res1 == 0:
                error1 = "Profile not shortlisted"
                return render_template("index.html", error1=error1 , mat=words)
            else:
                return render_template("index.html", matches=matches, res1=res1, res2=res2)

    return jsonify(error='Invalid file or words input.')

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in {'pdf', 'docx'}

def extract_text_from_pdf(file):
    pdf_document = fitz.open(stream=file.stream.read(), filetype='pdf')
    text = ''
    for page in pdf_document:
        text += page.get_text()
    pdf_document.close()
    return text.lower()

def extract_text_from_docx(file):
    doc = Document(file)
    full_text = []
    for paragraph in doc.paragraphs:
        full_text.append(paragraph.text)
    return ' '.join(full_text).lower()

def store_matches(matches,name):
    create_table()
    cur = mysql.connection.cursor()
    
    for i in matches:
        k=str(i)
        cur.execute('INSERT INTO skills_list(name, skills) VALUES (%s, %s)', (name, k))
    mysql.connection.commit()
    cur.close()

def store_candidateDetails(name, email, mobile_number):
    create_candidatesTable()
    file = request.files['file']
    temp_path = f"temp/{file.filename}"
    file.save(temp_path)
    with open(temp_path, 'rb') as f:
      file_content   = f.read()
    cur = mysql.connection.cursor()
    cur.execute('INSERT INTO candidates (name, email, mobile_number, resume) VALUES (%s, %s, %s, %s)', (name, email, mobile_number, file_content))
    mysql.connection.commit()
    cur.close()
if __name__ == '__main__':
    app.run(debug=True)
